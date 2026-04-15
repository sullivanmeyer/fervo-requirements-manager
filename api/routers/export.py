"""Requirements document export — Word (.docx) and PDF formats.

Generates a structured document organized by the system hierarchy tree, with
requirements listed under each node.  Accepts the same filter parameters as
the requirements list endpoint so the export scope matches what the user sees
in the table.

Word format uses python-docx (headings + paragraphs).
PDF format uses reportlab Platypus (headings + paragraphs).

Both formats are streamed back as file downloads so the browser triggers Save As.
"""
from __future__ import annotations

import io
from typing import Any, Optional
from uuid import UUID

from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from database import get_db
from models import HierarchyNode, Requirement

router = APIRouter()

SELF_DERIVED_ID = "SELF-000"


# ---------------------------------------------------------------------------
# Hierarchy helpers
# ---------------------------------------------------------------------------

def _build_tree_index(db: Session) -> tuple[dict[str, HierarchyNode], dict[str, list[str]]]:
    """Return (node_by_id, children_map) for all non-archived nodes."""
    all_nodes = (
        db.query(HierarchyNode)
        .filter(HierarchyNode.archived == False)  # noqa: E712
        .order_by(HierarchyNode.sort_order, HierarchyNode.name)
        .all()
    )
    node_by_id: dict[str, HierarchyNode] = {str(n.id): n for n in all_nodes}
    children_map: dict[str, list[str]] = {str(n.id): [] for n in all_nodes}
    for n in all_nodes:
        if n.parent_id and str(n.parent_id) in children_map:
            children_map[str(n.parent_id)].append(str(n.id))
    return node_by_id, children_map


def _dfs_order(node_by_id: dict, children_map: dict) -> list[tuple[str, int]]:
    """Return (node_id, depth) for every node in depth-first tree order."""
    roots = [nid for nid, n in node_by_id.items() if n.parent_id is None]
    result: list[tuple[str, int]] = []

    def dfs(nid: str, depth: int) -> None:
        result.append((nid, depth))
        for child_id in children_map.get(nid, []):
            dfs(child_id, depth + 1)

    for root_id in sorted(roots, key=lambda nid: node_by_id[nid].sort_order):
        dfs(root_id, 0)
    return result


def _req_dict(req: Requirement, blocks: list[dict] | None = None) -> dict[str, Any]:
    return {
        "id": req.requirement_id,
        "title": req.title,
        "classification": req.classification,
        "classification_subtype": req.classification_subtype,
        "statement": req.statement,
        "rationale": req.rationale,
        "source_clause": req.source_clause,
        "owner": req.owner,
        "status": req.status,
        "discipline": req.discipline,
        "source_document": req.source_document.title if req.source_document else None,
        "stale": req.stale,
        "content_source": req.content_source,
        "blocks": blocks or [],
    }


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _parse_pipe_table(content: str) -> dict | None:
    """
    Parse newline-and-pipe-delimited block content into a table_data dict.

    Used as a fallback when a table_block has no structured table_data
    (e.g. blocks extracted by pdfplumber before Vision parsing was active).

    Each line becomes a row; `|` separates cells.  The first row is treated
    as the header row when there are multiple rows.  Returns None if the
    content has fewer than two pipe characters (not obviously tabular).
    """
    if content.count("|") < 2:
        return None
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    if not lines:
        return None
    rows = [[cell.strip() for cell in ln.split("|")] for ln in lines]
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    if len(rows) > 1:
        return {"caption": None, "headers": rows[0], "rows": rows[1:], "context_note": None}
    return {"caption": None, "headers": [], "rows": rows, "context_note": None}


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def _generate_word(
    doc_title: str,
    nodes_with_reqs: list[tuple[HierarchyNode, int, list[dict]]],
    unassigned: list[dict],
) -> bytes:
    from docx import Document  # type: ignore[import]
    from docx.shared import Pt, RGBColor  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]

    doc = Document()
    doc.core_properties.title = doc_title

    title_para = doc.add_heading(doc_title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _render_table_word(td: dict) -> None:
        """Render a table_data dict as a Word table."""
        caption = td.get("caption")
        raw_headers = td.get("headers") or []
        rows = td.get("rows") or []
        footnotes = td.get("footnotes")

        if not raw_headers and not rows:
            return

        # Normalize headers to list-of-rows
        if raw_headers and isinstance(raw_headers[0], str):
            header_rows: list[list[str]] = [raw_headers]  # type: ignore[list-item]
        else:
            header_rows = raw_headers  # type: ignore[assignment]

        all_rows = list(header_rows) + list(rows)
        n_cols = max((len(r) for r in all_rows), default=1)

        if caption:
            cp = doc.add_paragraph(caption)
            for r in cp.runs:
                r.font.size = Pt(9)
                r.italic = True

        n_total = len(header_rows) + len(rows)
        if n_total == 0:
            return

        table = doc.add_table(rows=n_total, cols=n_cols, style="Table Grid")

        for ri, hrow in enumerate(header_rows):
            cells = table.rows[ri].cells
            for ci in range(n_cols):
                val = hrow[ci] if ci < len(hrow) else ""
                cells[ci].text = str(val) if val is not None else ""
                if cells[ci].paragraphs[0].runs:
                    run = cells[ci].paragraphs[0].runs[0]
                    run.bold = True
                    run.font.size = Pt(9)

        for ri, drow in enumerate(rows):
            cells = table.rows[len(header_rows) + ri].cells
            for ci in range(n_cols):
                val = drow[ci] if ci < len(drow) else ""
                cells[ci].text = str(val) if val is not None else ""
                if cells[ci].paragraphs[0].runs:
                    cells[ci].paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph("")  # spacer after table

        if footnotes:
            fn = doc.add_paragraph(footnotes)
            for r in fn.runs:
                r.font.size = Pt(8)
                r.italic = True

    def _render_block_word(block: dict) -> None:
        bt = block.get("block_type", "")
        td = block.get("table_data") or (
            _parse_pipe_table(block.get("content", "")) if bt == "table_block" else None
        )
        content = block.get("content", "")

        if bt == "table_block" and td:
            _render_table_word(td)
        elif bt == "heading":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(10)
        else:
            doc.add_paragraph(content)

    def _add_requirement(req: dict) -> None:
        p = doc.add_paragraph()
        run = p.add_run(f"[{req['id']}]  {req['title']}")
        run.bold = True
        if req.get("stale"):
            run.font.color.rgb = RGBColor(0xB4, 0x5A, 0x09)

        class_text = req["classification"]
        if req.get("classification_subtype"):
            class_text += f" — {req['classification_subtype']}"
        stale_note = "  ⚠ STALE" if req.get("stale") else ""
        meta = doc.add_paragraph(
            f"Classification: {class_text}  |  Discipline: {req['discipline']}"
            f"  |  Status: {req['status']}  |  Owner: {req['owner']}{stale_note}"
        )
        for run in meta.runs:
            run.font.size = Pt(9)
        meta.paragraph_format.space_after = Pt(2)

        if req.get("content_source") == "block_linked" and req.get("blocks"):
            for block in req["blocks"]:
                _render_block_word(block)
        else:
            doc.add_paragraph(req["statement"])

        source_parts = []
        if req.get("source_document"):
            source_parts.append(f"Source: {req['source_document']}")
        if req.get("source_clause"):
            source_parts.append(f"Clause: {req['source_clause']}")
        if source_parts:
            sp = doc.add_paragraph("  ".join(source_parts))
            for run in sp.runs:
                run.font.size = Pt(9)

        doc.add_paragraph("")  # spacer

    for node, depth, reqs in nodes_with_reqs:
        level = min(depth + 1, 4)
        doc.add_heading(node.name, level)
        for req in reqs:
            _add_requirement(req)

    if unassigned:
        doc.add_heading("Unassigned Requirements", 1)
        for req in unassigned:
            _add_requirement(req)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_pdf(
    doc_title: str,
    nodes_with_reqs: list[tuple[HierarchyNode, int, list[dict]]],
    unassigned: list[dict],
) -> bytes:
    from reportlab.lib.pagesizes import LETTER  # type: ignore[import]
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore[import]
    from reportlab.lib.units import inch  # type: ignore[import]
    from reportlab.lib import colors  # type: ignore[import]
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore[import]

    buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=doc_title,
    )
    styles = getSampleStyleSheet()

    h_styles = [
        ParagraphStyle("H1", parent=styles["Heading1"], fontSize=16, spaceAfter=8),
        ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, spaceAfter=6),
        ParagraphStyle("H3", parent=styles["Heading3"], fontSize=12, spaceAfter=4),
        ParagraphStyle("H4", parent=styles["Heading4"], fontSize=11, spaceAfter=4),
    ]
    req_id_style = ParagraphStyle(
        "ReqId", parent=styles["Normal"], fontSize=10,
        fontName="Helvetica-Bold", spaceAfter=2,
    )
    req_id_stale_style = ParagraphStyle(
        "ReqIdStale", parent=styles["Normal"], fontSize=10,
        fontName="Helvetica-Bold", textColor=colors.HexColor("#B45A09"), spaceAfter=2,
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["Normal"], fontSize=8,
        textColor=colors.grey, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=10, spaceAfter=4,
    )

    story = []
    story.append(Paragraph(doc_title, styles["Title"]))
    story.append(Spacer(1, 0.2 * inch))

    # Page width minus 1-inch margins on each side
    CONTENT_WIDTH = 6.5 * inch

    def _render_table_pdf(td: dict) -> None:
        """Render a table_data dict as a ReportLab Table flowable."""
        from reportlab.platypus import Table as RLTable, TableStyle  # type: ignore[import]

        caption = td.get("caption")
        raw_headers = td.get("headers") or []
        rows = td.get("rows") or []
        footnotes = td.get("footnotes")

        if not raw_headers and not rows:
            return

        # Normalize headers to list-of-rows
        if raw_headers and isinstance(raw_headers[0], str):
            header_rows: list[list] = [list(raw_headers)]
        else:
            header_rows = [list(r) for r in raw_headers]

        all_rows = header_rows + [list(r) for r in rows]
        n_cols = max((len(r) for r in all_rows), default=1)

        # Pad every row to n_cols
        padded = [r + [""] * (n_cols - len(r)) for r in all_rows]

        if caption:
            story.append(Paragraph(caption, meta_style))

        col_widths = [CONTENT_WIDTH / n_cols] * n_cols
        tbl = RLTable(padded, colWidths=col_widths, hAlign="LEFT")

        n_header_rows = len(header_rows)
        tbl.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, n_header_rows - 1), "Helvetica-Bold"),
            ("BACKGROUND", (0, 0), (-1, n_header_rows - 1), colors.HexColor("#E8E8E8")),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("WORDWRAP", (0, 0), (-1, -1), True),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.1 * inch))

        if footnotes:
            story.append(Paragraph(footnotes, meta_style))

    def _render_block_pdf(block: dict) -> None:
        bt = block.get("block_type", "")
        td = block.get("table_data") or (
            _parse_pipe_table(block.get("content", "")) if bt == "table_block" else None
        )
        content = block.get("content", "")

        if bt == "table_block" and td:
            _render_table_pdf(td)
        elif bt == "heading":
            heading_style = ParagraphStyle(
                "BlockHeading", parent=styles["Normal"],
                fontSize=10, fontName="Helvetica-Bold", spaceAfter=2,
            )
            story.append(Paragraph(content, heading_style))
        else:
            story.append(Paragraph(content, body_style))

    def _add_requirement(req: dict) -> None:
        id_text = f"[{req['id']}]  {req['title']}"
        if req.get("stale"):
            id_text += "  ⚠ STALE"
            story.append(Paragraph(id_text, req_id_stale_style))
        else:
            story.append(Paragraph(id_text, req_id_style))

        class_text = req["classification"]
        if req.get("classification_subtype"):
            class_text += f" — {req['classification_subtype']}"
        meta = (
            f"Classification: {class_text} | Discipline: {req['discipline']} "
            f"| Status: {req['status']} | Owner: {req['owner']}"
        )
        story.append(Paragraph(meta, meta_style))

        if req.get("content_source") == "block_linked" and req.get("blocks"):
            for block in req["blocks"]:
                _render_block_pdf(block)
        else:
            story.append(Paragraph(req["statement"], body_style))

        source_parts = []
        if req.get("source_document"):
            source_parts.append(f"Source: {req['source_document']}")
        if req.get("source_clause"):
            source_parts.append(f"Clause: {req['source_clause']}")
        if source_parts:
            story.append(Paragraph("  ".join(source_parts), meta_style))

        story.append(Spacer(1, 0.1 * inch))

    for node, depth, reqs in nodes_with_reqs:
        level = min(depth, 3)
        story.append(Paragraph(node.name, h_styles[level]))
        for req in reqs:
            _add_requirement(req)

    if unassigned:
        story.append(Paragraph("Unassigned Requirements", h_styles[0]))
        for req in unassigned:
            _add_requirement(req)

    pdf_doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Export endpoint
# ---------------------------------------------------------------------------

@router.get("/export/requirements-document")
def export_requirements(
    format: str = Query("word", pattern="^(word|pdf)$"),
    doc_title: str = Query("Requirements Document"),
    status: Optional[list[str]] = Query(None),
    classification: Optional[str] = Query(None),
    classification_subtype: Optional[str] = Query(None),
    discipline: Optional[list[str]] = Query(None),
    owner: Optional[str] = Query(None),
    hierarchy_node_id: Optional[str] = Query(None),
    include_descendants: bool = Query(False),
    stale: Optional[bool] = Query(None),
    db: Session = Depends(get_db),
):
    """
    Export all matching requirements as a Word (.docx) or PDF file, organized
    by the system hierarchy tree.  Accepts a subset of the list-requirements
    filter params so the exported scope matches what the user sees in the table.

    Requirements are grouped under their hierarchy node headings in tree order.
    Requirements assigned to no node appear at the end under "Unassigned".
    """
    from models import HierarchyNode as HN
    from sqlalchemy import or_

    base_q = db.query(Requirement).filter(
        Requirement.requirement_id != SELF_DERIVED_ID
    )
    if status:
        base_q = base_q.filter(Requirement.status.in_(status))
    if classification:
        base_q = base_q.filter(Requirement.classification == classification)
    if classification_subtype:
        base_q = base_q.filter(Requirement.classification_subtype == classification_subtype)
    if discipline:
        base_q = base_q.filter(Requirement.discipline.in_(discipline))
    if owner:
        base_q = base_q.filter(Requirement.owner.ilike(f"%{owner}%"))
    if stale is not None:
        base_q = base_q.filter(Requirement.stale == stale)

    if hierarchy_node_id:
        from routers.requirements import _collect_descendant_ids
        if include_descendants:
            node_ids = _collect_descendant_ids(hierarchy_node_id, db)
        else:
            node_ids = {UUID(hierarchy_node_id)}
        base_q = base_q.filter(
            Requirement.hierarchy_nodes.any(HN.id.in_(node_ids))
        )

    reqs = base_q.order_by(Requirement.requirement_id).all()

    # Batch-load linked blocks for block_linked requirements (single query)
    from models import RequirementBlock, DocumentBlock as DocBlock
    blocks_by_req: dict[str, list[dict]] = {}
    block_linked_ids = [req.id for req in reqs if req.content_source == "block_linked"]
    if block_linked_ids:
        junctions = (
            db.query(RequirementBlock, DocBlock)
            .join(DocBlock, RequirementBlock.block_id == DocBlock.id)
            .filter(RequirementBlock.requirement_id.in_(block_linked_ids))
            .order_by(RequirementBlock.requirement_id, RequirementBlock.sort_order)
            .all()
        )
        for junc, blk in junctions:
            rid = str(junc.requirement_id)
            if rid not in blocks_by_req:
                blocks_by_req[rid] = []
            blocks_by_req[rid].append({
                "block_type": blk.block_type,
                "content": blk.content,
                "clause_number": blk.clause_number,
                "heading_text": blk.heading,
                "table_data": blk.table_data,
            })

    # Pre-build all req dicts (with blocks injected for block_linked reqs)
    req_dicts: dict[str, dict] = {
        str(req.id): _req_dict(req, blocks_by_req.get(str(req.id)))
        for req in reqs
    }

    # Build hierarchy index
    node_by_id, children_map = _build_tree_index(db)
    ordered_nodes = _dfs_order(node_by_id, children_map)

    # Group requirements by hierarchy node id
    node_to_reqs: dict[str, list[dict]] = {nid: [] for nid in node_by_id}
    assigned_req_ids: set[str] = set()
    for req in reqs:
        for node in req.hierarchy_nodes:
            nid = str(node.id)
            if nid in node_to_reqs:
                node_to_reqs[nid].append(req_dicts[str(req.id)])
                assigned_req_ids.add(str(req.id))

    # Build the ordered list of (node, depth, reqs) — skip empty nodes
    nodes_with_reqs: list[tuple[HierarchyNode, int, list[dict]]] = []
    for nid, depth in ordered_nodes:
        if node_to_reqs.get(nid):
            nodes_with_reqs.append((node_by_id[nid], depth, node_to_reqs[nid]))

    # Requirements not assigned to any hierarchy node
    unassigned = [req_dicts[str(r.id)] for r in reqs if str(r.id) not in assigned_req_ids]

    if format == "word":
        content = _generate_word(doc_title, nodes_with_reqs, unassigned)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{doc_title.replace(' ', '_')}.docx"
    else:
        content = _generate_pdf(doc_title, nodes_with_reqs, unassigned)
        media_type = "application/pdf"
        filename = f"{doc_title.replace(' ', '_')}.pdf"

    return StreamingResponse(
        io.BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
